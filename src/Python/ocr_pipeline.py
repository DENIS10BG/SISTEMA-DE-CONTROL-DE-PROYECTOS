"""
================================================================================
🚀 SISTEMA INTELIGENTE DE INTEGRIDAD DOCUMENTAL - PIPELINE OCR OPTIMIZADO (v3.2)
================================================================================
Mejoras finales:
  - Ajustes de rendimiento: DPI=200, limit_side_len=1440, workers=4, lote=30
  - Detección robusta de tablas horizontales y verticales
  - Filtrado agresivo de líneas basura (números, símbolos)
  - Extracción de títulos (líneas anteriores a la tabla)
  - Paralelismo real y memoria controlada
  - NUEVO: Modo completo en primeras 15 páginas para preservar tablas (índice, inventario)
  - CORREGIDO: Guardado de resultados.json con páginas ordenadas
================================================================================
"""

import os
import json
import base64
import subprocess
import threading
import logging
import re
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Dict, List, Optional

import cv2
import numpy as np
from pdf2image import convert_from_path, pdfinfo_from_path
from tqdm import tqdm
from docx import Document

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s", datefmt="%H:%M:%S")
logger = logging.getLogger("OCR_Optimizado")


class Config:
    # --- RUTAS (AJUSTA SEGÚN TU SISTEMA) ---
    pdf_path: Path = Path(r"C:\Users\denis\Downloads\proyecto_ocr_2.pdf")
    output_dir: Path = Path(r"C:\Users\denis\Downloads\Resultado_Final")
    paddle_base: Path = Path(r"C:\Users\denis\Downloads\PaddleOCR-json_v1.4.1_windows_x64\PaddleOCR-json_v1.4.1")
    poppler_path: Path = Path(r"C:\Users\denis\Downloads\Release-26.02.0-0\poppler-26.02.0\Library\bin")
    ocr_model: str = "models/config_en.txt"   # Puedes probar con "config_en_ppocrv3.txt" si existe

    # --- PARÁMETROS DE RENDIMIENTO (OPTIMIZADOS) ---
    dpi: int = 105                    # 150-200 es suficiente para OCR, más rápido que 300
    limit_side_len: int = 950         # Reduce el tamaño de imagen, acelera OCR
    num_workers: int = 4               # Núcleos FÍSICOS de tu CPU (no hilos)
    paginas_por_lote: int = 30         # Menos páginas en RAM, evita swapping
    timeout_por_pagina: int = 25      # Tiempo máximo por página
    # --- NUEVO: Páginas en modo completo (preservar tablas) ---
    paginas_modo_completo: int = 10   # Las primeras 15 páginas conservan tablas


class MotorVision:
    @staticmethod
    def pil_a_cv2(imagen_pil) -> np.ndarray:
        return cv2.cvtColor(np.array(imagen_pil), cv2.COLOR_RGB2BGR)

    @staticmethod
    def procesar_para_ocr(img: np.ndarray) -> np.ndarray:
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        return clahe.apply(gray)


class MotorOCRPersistente:
    """Lanza PaddleOCR-json una sola vez y reutiliza el proceso."""
    def __init__(self, config: "Config", worker_id: int):
        self.config = config
        self.worker_id = worker_id
        self.lock = threading.Lock()
        exe = config.paddle_base / "PaddleOCR-json.exe"

        self.proc = subprocess.Popen(
            [str(exe), f"-config_path={config.ocr_model}", f"-limit_side_len={config.limit_side_len}"],
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            stderr=subprocess.DEVNULL,
            cwd=str(config.paddle_base),
            bufsize=1,
        )
        self._esperar_arranque()
        logger.info(f"Worker {worker_id}: motor OCR listo (PID {self.proc.pid}).")

    def _esperar_arranque(self, max_lineas: int = 200):
        for _ in range(max_lineas):
            linea = self.proc.stdout.readline()
            if not linea:
                raise RuntimeError(f"Worker {self.worker_id}: proceso OCR murió en arranque.")
            texto = linea.decode("utf-8", errors="ignore").strip()
            if texto:
                logger.debug(f"[worker {self.worker_id} init] {texto}")
            if "init completed" in texto.lower():
                return
        logger.warning(f"Worker {self.worker_id}: no se detectó 'init completed', se continúa.")

    def procesar_base64(self, imagen_b64: str) -> Dict:
        comando = json.dumps({"image_base64": imagen_b64}, ensure_ascii=True, indent=None) + "\n"
        with self.lock:
            try:
                self.proc.stdin.write(comando.encode("utf-8"))
                self.proc.stdin.flush()
                linea = self.proc.stdout.readline()
                if not linea:
                    return {"code": 500, "data": "proceso sin respuesta"}
                return json.loads(linea.decode("utf-8", errors="ignore"))
            except Exception as e:
                return {"code": 500, "data": str(e)}

    def cerrar(self):
        try:
            self.proc.stdin.close()
        except Exception:
            pass
        try:
            self.proc.terminate()
            self.proc.wait(timeout=5)
        except Exception:
            try:
                self.proc.kill()
            except Exception:
                pass


class AnalizadorEstructura:
    """Procesa los resultados del OCR y extrae títulos de tablas (modo normal)
       o todo el texto (modo completo)."""
    @staticmethod
    def procesar_resultado(resultado: Dict, modo_completo: bool = False) -> str:
        if resultado.get("code") != 100:
            return ""

        items = []
        for item in resultado.get("data", []):
            box = item.get("box", [])
            text = item.get("text", "").strip()

            if "escaneado con camscanner" in text.lower():
                continue

            if box and len(box) >= 4 and text:
                y_center = (box[0][1] + box[2][1]) / 2
                altura = box[2][1] - box[0][1]
                x_inicio = box[0][0]
                x_fin = box[2][0]

                if altura > 8:
                    items.append({
                        "text": text, "y": y_center,
                        "x_inicio": x_inicio, "x_fin": x_fin, "altura": altura
                    })

        if not items:
            return ""

        if modo_completo:
            return AnalizadorEstructura._reconstruir_texto_completo(items)
        else:
            return AnalizadorEstructura._reconstruir_geometria(items)

    @staticmethod
    def _reconstruir_texto_completo(items: List[Dict]) -> str:
        """Reconstruye todo el texto de la página sin filtrar tablas."""
        if not items:
            return ""

        items.sort(key=lambda i: i["y"])
        lineas = []
        linea_actual = []
        y_referencia = items[0]["y"]
        umbral_y = items[0]["altura"] * 0.7

        for item in items:
            if abs(item["y"] - y_referencia) <= umbral_y:
                linea_actual.append(item)
            else:
                lineas.append(linea_actual)
                linea_actual = [item]
                y_referencia = item["y"]
                umbral_y = item["altura"] * 0.7
        if linea_actual:
            lineas.append(linea_actual)

        texto_final = []
        for fila in lineas:
            fila.sort(key=lambda i: i["x_inicio"])
            renglon = ""
            x_anterior_fin = 0
            for palabra in fila:
                distancia_x = palabra["x_inicio"] - x_anterior_fin
                if x_anterior_fin != 0 and distancia_x > 50:
                    renglon += "   |   " + palabra["text"]
                else:
                    renglon += " " + palabra["text"] if x_anterior_fin != 0 else palabra["text"]
                x_anterior_fin = palabra["x_fin"]
            texto_final.append(renglon.strip())

        return "\n".join(texto_final)

    @staticmethod
    def _es_linea_tabla_horizontal(fila_items: List[Dict]) -> bool:
        """Detecta si una fila pertenece a una tabla horizontal (muchas columnas)."""
        if len(fila_items) < 4:
            return False
        ordenados = sorted(fila_items, key=lambda i: i["x_inicio"])
        separaciones = []
        for i in range(1, len(ordenados)):
            espacio = ordenados[i]["x_inicio"] - ordenados[i-1]["x_fin"]
            separaciones.append(espacio)
        if len(separaciones) >= 3:
            promedio_sep = sum(separaciones) / len(separaciones)
            if promedio_sep > 60:
                return True
        return False

    @staticmethod
    def _es_linea_valida(linea: str) -> bool:
        """Filtra líneas que son puramente numéricas, simbólicas o basura."""
        if not linea.strip():
            return False
        alfanumericos = sum(c.isalnum() or c.isspace() for c in linea)
        if alfanumericos / len(linea) < 0.2:
            return False
        palabras = linea.split()
        if len(palabras) >= 4:
            num_numericos = sum(1 for p in palabras if re.match(r'^[\d.,]+$', p))
            if num_numericos / len(palabras) > 0.5:
                return False
        simbolos = sum(1 for c in linea if c in '|-_=+*')
        if simbolos / len(linea) > 0.1:
            return False
        return True

    @staticmethod
    def _reconstruir_geometria(items: List[Dict]) -> str:
        """Reconstruye solo títulos de tablas (comportamiento original)."""
        if not items:
            return ""

        items.sort(key=lambda i: i["y"])
        lineas = []
        linea_actual = []
        y_referencia = items[0]["y"]
        umbral_y = items[0]["altura"] * 0.7

        for item in items:
            if abs(item["y"] - y_referencia) <= umbral_y:
                linea_actual.append(item)
            else:
                lineas.append(linea_actual)
                linea_actual = [item]
                y_referencia = item["y"]
                umbral_y = item["altura"] * 0.7
        if linea_actual:
            lineas.append(linea_actual)

        fila_metrics = []
        for fila in lineas:
            if not fila:
                continue
            num_palabras = len(fila)
            alturas = [p["altura"] for p in fila]
            altura_prom = sum(alturas) / len(alturas)
            fila_metrics.append({
                "num": num_palabras,
                "altura_prom": altura_prom,
                "fila": fila
            })

        umbral_num = 0.3
        umbral_alt = 0.3
        mejor_inicio = -1
        mejor_longitud = 0

        i = 0
        while i < len(fila_metrics):
            if fila_metrics[i]["num"] >= 3:
                j = i
                base_num = fila_metrics[i]["num"]
                base_alt = fila_metrics[i]["altura_prom"]
                while j < len(fila_metrics):
                    m = fila_metrics[j]
                    if (abs(m["num"] - base_num) / max(1, base_num) <= umbral_num and
                        abs(m["altura_prom"] - base_alt) / max(1, base_alt) <= umbral_alt):
                        j += 1
                    else:
                        break
                longitud = j - i
                if longitud >= 3 and longitud > mejor_longitud:
                    mejor_longitud = longitud
                    mejor_inicio = i
                i = j
            else:
                i += 1

        if mejor_inicio != -1:
            filas_titulo = []
            for idx in range(0, mejor_inicio):
                fila = fila_metrics[idx]["fila"]
                fila.sort(key=lambda p: p["x_inicio"])
                texto_fila = " ".join([p["text"] for p in fila])
                if AnalizadorEstructura._es_linea_valida(texto_fila):
                    filas_titulo.append(texto_fila)
            if filas_titulo:
                return "\n".join(filas_titulo)
            else:
                return ""

        tabla_horizontal_detectada = False
        primer_idx = 0
        for idx in range(len(fila_metrics) - 2):
            if (AnalizadorEstructura._es_linea_tabla_horizontal(fila_metrics[idx]["fila"]) and
                AnalizadorEstructura._es_linea_tabla_horizontal(fila_metrics[idx+1]["fila"]) and
                AnalizadorEstructura._es_linea_tabla_horizontal(fila_metrics[idx+2]["fila"])):
                tabla_horizontal_detectada = True
                primer_idx = idx
                break

        if tabla_horizontal_detectada:
            filas_titulo = []
            for idx in range(0, primer_idx):
                fila = fila_metrics[idx]["fila"]
                fila.sort(key=lambda p: p["x_inicio"])
                texto_fila = " ".join([p["text"] for p in fila])
                if AnalizadorEstructura._es_linea_valida(texto_fila):
                    filas_titulo.append(texto_fila)
            if filas_titulo:
                return "\n".join(filas_titulo)
            else:
                return ""

        texto_final = []
        for fila in lineas:
            fila.sort(key=lambda i: i["x_inicio"])
            renglon = ""
            x_anterior_fin = 0
            for palabra in fila:
                distancia_x = palabra["x_inicio"] - x_anterior_fin
                if x_anterior_fin != 0 and distancia_x > 50:
                    renglon += "   |   " + palabra["text"]
                else:
                    renglon += " " + palabra["text"] if x_anterior_fin != 0 else palabra["text"]
                x_anterior_fin = palabra["x_fin"]
            linea_texto = renglon.strip()
            if AnalizadorEstructura._es_linea_valida(linea_texto):
                texto_final.append(linea_texto)

        return "\n".join(texto_final)


class Orquestador:
    def __init__(self):
        self.config = Config()
        self.vision = MotorVision()
        self.config.output_dir.mkdir(parents=True, exist_ok=True)

        logger.info(f"Arrancando {self.config.num_workers} motores OCR persistentes...")
        self.motores: List[MotorOCRPersistente] = [
            MotorOCRPersistente(self.config, i) for i in range(self.config.num_workers)
        ]

    def _procesar_pagina(self, indice: int, img_pil) -> tuple:
        motor = self.motores[indice % self.config.num_workers]

        img_cv2 = self.vision.pil_a_cv2(img_pil)
        img_procesada = self.vision.procesar_para_ocr(img_cv2)

        ok, buffer = cv2.imencode(".png", img_procesada)
        if not ok:
            return indice, ""
        imagen_b64 = base64.b64encode(buffer).decode("ascii")

        resultado = motor.procesar_base64(imagen_b64)

        # Decidir modo: primeras N páginas conservan tablas (modo completo)
        modo_completo = indice <= self.config.paginas_modo_completo
        texto = AnalizadorEstructura.procesar_resultado(resultado, modo_completo=modo_completo)
        return indice, texto

    def ejecutar(self):
        logger.info("Iniciando Pipeline Optimizado...")

        info = pdfinfo_from_path(str(self.config.pdf_path), poppler_path=str(self.config.poppler_path))
        total_paginas = info["Pages"]
        logger.info(f"Documento con {total_paginas} páginas. Lotes de {self.config.paginas_por_lote} páginas.")
        logger.info(f"Modo completo (con tablas) activado para las primeras {self.config.paginas_modo_completo} páginas.")

        resultados: Dict[int, str] = {}
        barra = tqdm(total=total_paginas, desc="Procesando PDF")

        inicio = 1
        try:
            while inicio <= total_paginas:
                fin = min(inicio + self.config.paginas_por_lote - 1, total_paginas)

                imagenes = convert_from_path(
                    str(self.config.pdf_path),
                    dpi=self.config.dpi,
                    poppler_path=str(self.config.poppler_path),
                    first_page=inicio,
                    last_page=fin,
                    thread_count=4,
                )

                with ThreadPoolExecutor(max_workers=self.config.num_workers) as executor:
                    futuros = {
                        executor.submit(self._procesar_pagina, inicio + offset, img): inicio + offset
                        for offset, img in enumerate(imagenes)
                    }
                    for futuro in as_completed(futuros):
                        idx = futuros[futuro]
                        try:
                            idx, texto = futuro.result(timeout=self.config.timeout_por_pagina)
                        except Exception as e:
                            logger.warning(f"Página {idx} falló: {e}")
                            texto = ""
                        resultados[idx] = texto
                        barra.update(1)

                del imagenes
                inicio = fin + 1
        finally:
            barra.close()
            for motor in self.motores:
                motor.cerrar()

        # --- GENERACIÓN DE SALIDAS ---
        doc = Document()
        doc.add_heading("EXTRACCIÓN ESTRUCTURAL OPTIMIZADA", 0)
        for i in range(1, total_paginas + 1):
            doc.add_heading(f"PÁGINA {i}", level=1)
            for parrafo in resultados.get(i, "").split("\n"):
                if parrafo.strip():
                    doc.add_paragraph(parrafo.strip())
            doc.add_page_break()

        salida_docx = self.config.output_dir / "Documento_Final.docx"
        doc.save(str(salida_docx))
        logger.info(f"DOCX guardado en: {salida_docx}")

        # ---- CORRECCIÓN: ORDENAR EL DICCIONARIO ANTES DE GUARDAR ----
        resultados_ordenados = {k: resultados[k] for k in sorted(resultados.keys())}

        salida_json = self.config.output_dir / "resultados.json"
        with open(salida_json, "w", encoding="utf-8") as f:
            json.dump(resultados_ordenados, f, ensure_ascii=False, indent=2)
        logger.info(f"JSON guardado en: {salida_json} (páginas ordenadas)")

        salida_md = self.config.output_dir / "resultado.md"
        with open(salida_md, "w", encoding="utf-8") as f:
            f.write("# EXTRACCIÓN ESTRUCTURAL OPTIMIZADA\n\n")
            for i in range(1, total_paginas + 1):
                f.write(f"## PÁGINA {i}\n\n")
                texto_pagina = resultados.get(i, "").strip()
                if texto_pagina:
                    f.write("\n\n".join(texto_pagina.split("\n")))
                    f.write("\n\n")
                else:
                    f.write("*[Página sin texto detectable]*\n\n")
        logger.info(f"Markdown guardado en: {salida_md}")

        logger.info("¡Proceso completado con éxito!")


if __name__ == "__main__":
    app = Orquestador()
    app.ejecutar()